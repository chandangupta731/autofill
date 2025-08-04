# -*- coding: utf-8 -*-
from flask import Flask, render_template, request, jsonify
import os
import json
import re

# --- Docx Handling ---
try:
    import docx # Needed for Word processing
except ImportError:
    print("! ERROR: python-docx library not found. Install using: pip install python-docx")
    docx = None

# --- Image Processing Dependencies (Optional, checked before use) ---
try:
    from paddleocr import PaddleOCR
except ImportError:
    print("! WARNING: paddleocr or paddlepaddle not found. Image processing will fail.")
    PaddleOCR = None

# --- Gemini Dependency (Optional, checked before use) ---
try:
    import google.generativeai as genai
except ImportError:
     print("! WARNING: google-generativeai not found. Gemini analysis will fail.")
     genai = None

import traceback
from werkzeug.utils import secure_filename
from pymongo import MongoClient, ReturnDocument
import pymongo.errors
from datetime import datetime
from bson import ObjectId
from flask_cors import CORS

app = Flask(__name__)

# --- Configuration ---
UPLOAD_FOLDER = os.path.join(os.getcwd(), 'uploads')
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'docx'}
# ***** IMPORTANT: Use environment variables for sensitive keys in production *****
API_KEY = os.environ.get("GEMINI_API_KEY") # Replace placeholder if needed
MONGO_URI = os.environ.get("MONGO_URI")
MONGO_DB_NAME = "id_docs_db"
MONGO_USERS_COLLECTION_NAME = "users"
MONGO_DOCUMENTS_COLLECTION_NAME = "documents"
GEMINI_MODEL_NAME = 'gemini-1.5-pro-latest' # Needed for image analysis
PASSPORT_REDACTION_PLACEHOLDER = "[PASSPORT NUMBER REDACTED]"
AADHAAR_REDACTION_PLACEHOLDER = "[AADHAAR NUMBER REDACTED]"
ID_NUMBER_NOT_FOUND = "NA"
WORD_DOC_TYPE = "word_form" # Special type for Word Q&A data

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Mapping from questions in the Word doc to keys used in the database/API
# (QUESTION_TO_KEY_MAP remains the same as before)
QUESTION_TO_KEY_MAP = {
    "What is your National ID number (e.g., Aadhaar card)": "national_id",
    "Marital Status": "marital_status",
    "Passport Number": "passport_number",
    "Home Address Street Address (Line 1)": "home_address_line1",
    "Home Address Street Address (Line 2)": "home_address_line2",
    "Home Address City": "home_address_city",
    "Home Address State": "home_address_state",
    "Home Address Zip Code": "home_address_zip",
    "Is your mailing address the same as your home address?": "mailing_address_same_as_home",
    "Primary phone number": "primary_phone",
    "Secondary phone number": "secondary_phone",
    "Please provide a valid email address for verification.": "email",
    "What is your current employer name?": "current_employer_name",
    "Present Employer Address Street Address (Line 2)" : "present_employer_address_line2",
    "Present Employer Address Street Address (Line 1)": "present_employer_address_line1",
    "Briefly describe your duties and roles in the present company.": "current_employer_duties",
    "Present Employer Address City" : "present_employer_address_city",
    "Present Employer Address State":"present_employer_address_state",
    "Present Employer Address ZIP Code": "present_employer_address_zip_code",
    "Present Employer Phone Number":"present_employer_phone_number",
    "Briefly describe your duties and roles in the previous company.": "prevoius_employer_duties",
    "What is the start date of your present employment? (dd-mm-yyyy)": "current_employment_start_date",
    "Name of Institution": "institution_name",
    "Institution Address Street Address (Line 1)":"institution_address_line1",
    "Institution Address Street Address (Line 2)":"institution_address_line2",
    "Institution Address City":"institution_address_city",
    "Institution Address State":"institution_address_state",
    "Institution Address ZIP Code":"institution_address_zip_code",
    "Institution Address Country":"institution_address_country",
    "What are your educational qualification details?  Mention the highest qualification": "highest_qualification",
    "Start date of the education institutes (dd-mm-yyyy)": "education_institute_start_date",
    "End date of the education institutes   (dd-mm-yyyy)": "education_institute_end_date",
    "What is your intended date of arrival in the U.S.? (dd-mm-yyyy)": "us_arrival_date", # Target field for formatting
    "What is your intended date of departure from the U.S.? (dd-mm-yyyy)": "us_departure_date", # Target field for formatting
    "What are your arrival and departure cities?": "us_arrival_departure_cities",
    "Provide flight details if known (departure and arrival).": "us_flight_details",
    "List the locations you plan to visit in the U.S.": "us_visit_locations",
    "U.S. Stay Address Street Address (Line 1)":"us_stay_address_line1",
    "U.S. Stay Address Street Address (Line 2)":"us_stay_address_line2",
    "U.S. Stay Address City":"us_stay_address_city",
    "U.S. Stay Address State":"us_stay_address_state",
    "U.S. Stay Address ZIP Code":"us_stay_address_zip_code",
    "Who is paying for your trip? (Self, Other Person, Present Employer, Employer in the U.S.) If Other Specify Below": "trip_payer",
    "Surnames of Person Paying for Trip":"trip_payer_surname",
    "Given Names of Person Paying for Trip":"trip_payer_given_name",
    "Telephone Number of Person Paying for Trip":"trip_payer_telephone",
    "Email Address of Person Paying for Trip":"trip_payer_email",
    "Relationship to You (Person Paying for Trip)":"trip_payer_relationship",
    "Is the address of the party paying for your trip the same as your Home or Mailing Address? if no provide detail below.":"trip_payer_address_same_as_home",
    "Payer Address Street (Line 1)":"payer_address_street_line1",
    "Payer Address Street (Line 2)":"payer_address_street_line2",
    "Payer Address City":"payer_address_city",
    "Payer Address State/Province":"payer_address_state",
    "Payer Address Postal Zone/ZIP Code":"payer_address_zip_code",
    "Payer Address Country/Region ":"payer_address_country",
    "Are there other people traveling with you?\n(If yes then provide name and relationship with you.)": "travel_companions",
    "Are there other people traveling with you?(If yes then provide name and relationship with you.)": "travel_companions", # Fallback
    "Have you ever been to the U.S.?(If yes then provide the travel dates and the length of the stay in U.S.)": "previous_us_visits",
    "Have you ever been issued a U.S. Visa?\n(If yes then please provide the visa number)": "previous_us_visa_issued",
    "Have you ever been issued a U.S. Visa?(If yes then please provide the visa number)": "previous_us_visa_issued", # Fallback
    "Have you ever been refused a U.S. visa or admission to the U.S., or withdrawn your application for admission at the port of entry? (Yes/No) If yes then please provide the reason, date and at which embassy it was rejected.": "previous_us_visa_refusal",
    "If yes then please provide the reason, date and at which embassy it was rejected.": "previous_us_visa_refusal_details",
    "Has anyone ever filed an immigrant petition on your behalf with the United States Citizenship and Immigration Services? (Yes/No) If yes then please explain.": "immigrant_petition_filed",
    "Contact person or organization in the U.S.": "us_contact_name_or_org",
    "Surnames of the US contact?": "us_contact_surname",
    "Given names of the US contact?": "us_contact_given_name",
    "What is the relationship to you?": "us_contact_relationship",
    "U.S. Street Address (Line 1)": "us_contact_street_address_line1",
    "U.S. Street Address (Line 2)": "us_contact_street_address_line2",
    "U.S. City": "us_contact_city",
    "U.S. State": "us_contact_state",
    "U.S. ZIP Code": "us_contact_zip",
    "Provide the email address of the U.S. point of contact.": "us_contact_email",
    "Provide the phone number of the U.S. point of contact.": "us_contact_phone",
    "Do you have a social media presence? Please list IDs of all accounts.": "social_media_presence",
    "Have you ever lost a passport or had one stolen?(If yes, then provide lost passport number and where it was lost and when)": "lost_stolen_passport_details",
    "What is your Father's  surname?": "father_surname",
    "What is your Father's  given name?": "father_given_name",
    "Father date of birth (dd-mm-yyyy)": "father_dob",
    "Is your father in the U.S.?": "father_in_us",
    "What is your Mother's  surname?": "mother_surname",
    "What is your Mother's given name?": "mother_given_name",
    "Mother date of birth (dd-mm-yyyy)": "mother_dob",
    "Is your mother in the U.S.?": "mother_in_us",
    "Do you have immediate relatives in the U.S. (not including parents)?": "immediate_relatives_in_us",
    "Do you have any other relatives in the U.S.?": "other_relatives_in_us",
    "Spouse Surname": "spouse_surname",
    "Spouse Given Name": "spouse_given_name",
    "Provide your spouse's date of birth    (dd-mm-yyyy)": "spouse_dob",
    "What is your spouse's place of birth (city, country)?": "spouse_pob",
    "What is your spouse's address?               ( Same as Home Address, Same as Mailing Address, Same as U.S. Contact Address, Do Not Know, Other (Specify Address))": "spouse_address",
    "Have you attended any educational institutions at a secondary level or above?": "attended_secondary_education_or_above",
    "Do you belong to a clan or tribe?": "clan_or_tribe",
    "Are you a permanent resident of a country/region other than your country/region of origin (nationality). If yes then provide name of the country or region.": "other_permanent_residency",
    "Do you hold or have you held any nationality other than the one indicated above on nationality?": "other_nationalities",
    "List all languages you speak.": "languages_spoken",
    "Have you travelled to any countries/regions in the last 5 years? List them.": "countries_visited_last_5_years",
    "Have you contributed to any professional, social, or charitable organization? (Yes or No)": "organization_contributions",
    "Monthly net income post tax deduction": "monthly_income",
    "Were you previously employed?": "previously_employed",

     # --- Detailed Previous Employer 1 ---
    "1.Previous employers name": "prev_emp_1_name",
    "1.Previous Employers Street Address (Line 1)": "prev_emp_1_address_line1",
    "1.Previous Employers Street Address (Line 2)":"prev_emp_1_address_line2",
    "1.Previous Employers Address City":"prev_emp_1_address_city",
    "1.Previous Employers Address State":"prev_emp_1_address_state",
    "1.Previous Employers Address ZIP Code":"prev_emp_1_address_zip_code",
    "1.Previous Employers Address Country":"prev_emp_1_address_country",
    "1.Previous Employers Phone": "prev_emp_1_address_phone",
    "1. What was your job title?": "prev_emp_1_job_title", # Added leading "1. " for consistency
    "1. Provide the supervisors surname": "prev_emp_1_supervisor_surname",
    "1. Provide the supervisors given name": "prev_emp_1_supervisor_given_name",
    "1. Employment start date (dd-mm-yyyy)": "prev_emp_1_start_date",
    "1.Previous employers Phone number": "prev_emp_1_phone",
    "1. Employment end date (dd-mm-yyyy)": "prev_emp_1_end_date",
    "1. Briefly describe your duties and roles in the previous company.": "prev_emp_1_duties",

    # --- Detailed Previous Employer 2 ---
    "2.Previous employers name 2": "prev_emp_2_name", # Assuming "2" is part of the question
    "2.Previous employers phone number": "prev_emp_2_phone",
    "2.Previous employers address": "prev_emp_2_address", # Added leading "2. "
    "2. What was your job title?": "prev_emp_2_job_title", # Added leading "2. "
    "2. Provide the supervisors name.": "prev_emp_2_supervisor_name", # Added leading "2. "
    "2. Employment start date (dd-mm-yyyy)": "prev_emp_2_start_date",
    "2. Employment end date (dd-mm-yyyy)": "prev_emp_2_end_date",
    "2. Briefly describe your duties and roles in the previous company.": "prev_emp_2_duties",
}


# --- MongoDB Setup ---
# (MongoDB setup code remains the same)
mongo_client = None
db = None
users_collection = None
documents_collection = None
try:
    print(f"Connecting to MongoDB at {MONGO_URI}...")
    # Increased timeout for potentially slower connections
    mongo_client = MongoClient(MONGO_URI, serverSelectionTimeoutMS=10000)
    # The ismaster command is cheap and does not require auth.
    mongo_client.admin.command('ismaster')
    print(f"Connected. Using DB: '{MONGO_DB_NAME}'")
    db = mongo_client[MONGO_DB_NAME]
    users_collection = db[MONGO_USERS_COLLECTION_NAME]
    documents_collection = db[MONGO_DOCUMENTS_COLLECTION_NAME]
    print(f"Using collections: '{MONGO_USERS_COLLECTION_NAME}', '{MONGO_DOCUMENTS_COLLECTION_NAME}'")
except pymongo.errors.ConnectionFailure as e:
    print(f"FATAL: MongoDB connection failed: {e}")
    mongo_client = None # Ensure client is None if connection fails
except Exception as e:
    print(f"FATAL: MongoDB setup error: {e}")
    print(traceback.format_exc())
    mongo_client = None # Ensure client is None on other setup errors


# --- Helper Functions ---
# (All helper functions remain the same: allowed_file, extract_full_text_from_word,
# find_passport_no_in_text, extract_text_paddle, detect_document_type,
# find_and_redact_passport_no, find_and_redact_aadhaar_no, analyze_text_with_gemini,
# compute_unique_identifiers, extract_data_from_word, sanitize_mongodb_keys,
# split_name, format_date_for_form, format_date_ddmmyyyy, parse_place_of_birth)

def allowed_file(filename):
    """Checks if the file extension is allowed."""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_full_text_from_word(file_path):
    """Extracts all text content from a .docx file."""
    print(f"Attempting to extract full text from Word file: {file_path}")
    if docx is None:
        print("Error: python-docx library is not available.")
        return None
    full_text_lines = []
    try:
        document = docx.Document(file_path)
        # Extract text from paragraphs
        for paragraph in document.paragraphs:
            full_text_lines.append(paragraph.text)
        # Extract text from tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text_lines.append(cell.text)
        # Join non-empty lines and strip leading/trailing whitespace
        combined_text = "\n".join(filter(None, full_text_lines)).strip()
        print(f"Successfully extracted text. Length: {len(combined_text)}")
        return combined_text
    except Exception as e:
        print(f"Error extracting text from Word file '{file_path}': {e}")
        print(traceback.format_exc())
        return None

def find_passport_no_in_text(text_content):
    """Finds the first Indian passport number pattern in the given text."""
    if not text_content:
        print("Skipping passport number search: No text content provided.")
        return None
    # Pattern: Letter followed by 7 digits, allowing optional space after letter
    pattern = re.compile(r'\b([A-Z]\s?\d{7})\b', re.IGNORECASE)
    # Find all matches and sort by position
    matches = sorted(pattern.finditer(text_content), key=lambda m: m.start())
    if matches:
        # Get the first match, convert to uppercase, remove space
        found_passport_no = matches[0].group(1).upper().replace(" ", "")
        print(f"Found passport number pattern: {found_passport_no}")
        return found_passport_no
    else:
        print("Passport number pattern not found in the text.")
        return None

def extract_text_paddle(image_path, lang='en'):
    """Extracts text from an image file using PaddleOCR."""
    print(f"Attempting OCR on image: {image_path}")
    if PaddleOCR is None:
        print("Error: PaddleOCR library is not available.")
        return None
    if not os.path.exists(image_path):
        print(f"Error: Image file not found at path: {image_path}")
        return None
    try:
        # Initialize PaddleOCR (consider making this a global instance if performance is critical)
        # use_gpu=False is safer for general compatibility
        ocr_engine = PaddleOCR(use_angle_cls=True, lang=lang, use_gpu=False, show_log=False)
        # Perform OCR
        result = ocr_engine.ocr(image_path, cls=True)
        extracted_text = ""
        # Process results: Structure is [[lines]] where each line is [bbox, [text, confidence]]
        if result and result[0]: # Check if result is not empty and has the first list element
             # Safely extract text, checking types and lengths
            texts = [
                line[1][0] for line in result[0]
                if line and isinstance(line, list) and len(line) >= 2 and
                   isinstance(line[1], (tuple, list)) and len(line[1]) >= 2 and line[1][0]
            ]
            extracted_text = "\n".join(texts)

        if not extracted_text:
            print("Warning: OCR process completed but yielded no text.")
            return "" # Return empty string instead of None if OCR runs but finds nothing

        print("OCR extraction successful.")
        return extracted_text.strip()
    except Exception as e:
        print(f"Error during PaddleOCR processing for '{image_path}': {e}")
        print(traceback.format_exc())
        return None # Return None on critical OCR failure

def detect_document_type(text_content):
    """Attempts to detect if the text content is from an Indian Passport or Aadhaar card."""
    if not text_content:
        print("Cannot detect document type: No text content provided.")
        return "unknown"

    print("Detecting document type from text...")
    text_lower = text_content.lower()

    # Keywords for Passport and Aadhaar
    passport_keywords = ["passport", "republic of india", "issuing authority", "surname", "given name"]
    aadhaar_keywords = ["aadhaar", "uidai", "unique identification", "enrollment no", "government of india", "आधार"] # Added Hindi

    # Regular expressions for typical ID patterns
    # Passport MRZ-like pattern (Letter followed by 7 digits, optional space)
    passport_pattern = re.compile(r'\b([A-Z]\s?\d{7})\b', re.IGNORECASE)
    # Aadhaar pattern (12 digits, possibly with spaces)
    aadhaar_pattern = re.search(r'\b\d{4}\s?\d{4}\s?\d{4}\b', text_lower)

    # Scoring based on keywords and patterns
    passport_score = sum(keyword in text_lower for keyword in passport_keywords)
    passport_score += 2 if passport_pattern.search(text_content) else 0 # Use original case text for regex if needed

    aadhaar_score = sum(keyword in text_lower for keyword in aadhaar_keywords)
    aadhaar_score += 2 if aadhaar_pattern else 0

    print(f"Document Type Detection Scores - Passport: {passport_score}, Aadhaar: {aadhaar_score}")

    # Determine type based on scores (require a minimum score to reduce misidentification)
    if passport_score > aadhaar_score and passport_score >= 2:
        print("Detected type: Passport")
        return "passport"
    if aadhaar_score > passport_score and aadhaar_score >= 2:
        print("Detected type: Aadhaar")
        return "aadhaar"

    # Handle cases where scores are equal or low but a strong pattern exists
    if aadhaar_pattern and aadhaar_score >= 1 and passport_score == 0:
         print("Detected type: Aadhaar (based on pattern and keyword)")
         return "aadhaar"
    if passport_pattern.search(text_content) and passport_score >= 1 and aadhaar_score == 0:
         print("Detected type: Passport (based on pattern and keyword)")
         return "passport"

    print("Could not reliably determine document type from text.")
    return "unknown"

def find_and_redact_passport_no(text):
    """
    Finds the first passport number in text, prints the redacted part,
    and returns the processed number and the redacted text.
    Includes detailed debugging output if no match is found with the primary pattern.
    """
    print("\n--- Attempting to find and redact passport number ---")
    if not text or not isinstance(text, str):
        print("Input text is empty or not a string. Skipping redaction.")
        return None, text

    print(f"Input text (raw): '{text[:200]}{'...' if len(text) > 200 else ''}'") # Print a snippet for brevity
    print(f"Input text (repr for special chars): {repr(text[:200])}{'...' if len(text) > 200 else ''}")

    # Current pattern: A letter, an optional space, 7 digits, followed by a word boundary. Case-insensitive.
    # Leading \b was removed in a previous iteration to allow matching if stuck to preceding word characters.
    pattern = re.compile(r'([A-Z]\s?\d{7})\b', re.IGNORECASE)

    found_number = None
    redacted_text = text
    original_redacted_segment = None

    matches = sorted(pattern.finditer(text), key=lambda m: m.start())

    if matches:
        match = matches[0]
        original_redacted_segment = match.group(1) # group(1) is the content of ([A-Z]\s?\d{7})
        found_number = original_redacted_segment.upper().replace(" ", "")
        redacted_text = text[:match.start()] + PASSPORT_REDACTION_PLACEHOLDER + text[match.end():]

        print(f"SUCCESS: Found passport number for redaction: {found_number}")
        print(f"SUCCESS: Original segment that was redacted: '{original_redacted_segment}'")
    else:
        print(f"\nINFO: Main passport pattern r'([A-Z]\\s?\\d{{7}})\\b' (case-insensitive) did NOT find any matches in the provided text.")
        print("--- Initiating Debug Search for Potential Components ---")

        # Debug Step 1: Look for any 7-digit sequences
        debug_digits_pattern = re.compile(r'\d{7}')
        digit_matches = list(debug_digits_pattern.finditer(text))
        if digit_matches:
            print(f"DEBUG: Found {len(digit_matches)} sequence(s) of exactly 7 digits:")
            for i, m in enumerate(digit_matches):
                context_start = max(0, m.start() - 15)
                context_end = min(len(text), m.end() + 15)
                print(f"  - 7-digit match #{i+1}: '{m.group(0)}' at pos {m.start()}. Context: '{text[context_start:m.start()]}>>>{m.group(0)}<<<{text[m.end():context_end]}'")
        else:
            print("DEBUG: No sequences of 7 consecutive digits found anywhere in the text.")

        # Debug Step 2: Look for Letter + 7 Digits (no space, no boundaries)
        debug_letter_digits_pattern = re.compile(r'[A-Z]\d{7}', re.IGNORECASE)
        letter_digit_matches = list(debug_letter_digits_pattern.finditer(text))
        if letter_digit_matches:
            print(f"DEBUG: Found {len(letter_digit_matches)} sequence(s) of [Letter][7 Digits] (no space, no boundaries):")
            for i, m in enumerate(letter_digit_matches):
                context_start = max(0, m.start() - 15)
                context_end = min(len(text), m.end() + 15)
                print(f"  - L+7D match #{i+1}: '{m.group(0)}' at pos {m.start()}. Context: '{text[context_start:m.start()]}>>>{m.group(0)}<<<{text[m.end():context_end]}'")
        else:
            print("DEBUG: No sequences of [Letter][7 Digits] (no space, no boundaries) found.")

        # Debug Step 3: Look for Letter + optional Space + 7 Digits (no boundaries at all)
        debug_letter_space_digits_pattern = re.compile(r'[A-Z]\s?\d{7}', re.IGNORECASE)
        letter_space_digit_matches = list(debug_letter_space_digits_pattern.finditer(text))
        if letter_space_digit_matches:
            print(f"DEBUG: Found {len(letter_space_digit_matches)} sequence(s) of [Letter][opt.Space][7 Digits] (NO boundaries):")
            for i, m in enumerate(letter_space_digit_matches):
                context_start = max(0, m.start() - 15)
                context_end = min(len(text), m.end() + 15)
                print(f"  - L[s?]7D match #{i+1}: '{m.group(0)}' at pos {m.start()}. Context: '{text[context_start:m.start()]}>>>{m.group(0)}<<<{text[m.end():context_end]}'")
        else:
            print("DEBUG: No sequences of [Letter][opt.Space][7 Digits] (NO boundaries) found.")
        
        # Debug Step 4: Check if the *original strict pattern* (with both \b) would have matched anything
        # This helps understand if the previous iteration's change (removing leading \b) was relevant
        # or if the problem is more fundamental.
        original_strict_pattern = re.compile(r'\b([A-Z]\s?\d{7})\b', re.IGNORECASE)
        strict_matches = list(original_strict_pattern.finditer(text))
        if strict_matches:
            print(f"DEBUG: The original stricter pattern r'\\b([A-Z]\\s?\\d{{7}})\\b' WOULD have found {len(strict_matches)} match(es). This implies the context around the number might be tricky for word boundaries.")
        else:
            print(f"DEBUG: The original stricter pattern r'\\b([A-Z]\\s?\\d{{7}})\\b' also finds no matches.")

        print("--- End Debug Search ---")

    return found_number, redacted_text

def find_and_redact_aadhaar_no(text):
    """Finds the first valid Aadhaar number in text and returns the number and redacted text."""
    print("Attempting to find and redact Aadhaar number...")
    # Pattern for 12 digits, allowing optional spaces in between groups of 4
    pattern = re.compile(r'\b(\d{4}(?:\s?\d{4}){2})\b')
    found_number = None
    redacted_text = text
    match = pattern.search(text)
    if match:
        # Validate the length after removing spaces
        potential_number = "".join(match.group(1).split())
        if len(potential_number) == 12:
            found_number = potential_number
            # Replace only the first valid occurrence
            redacted_text = text[:match.start()] + AADHAAR_REDACTION_PLACEHOLDER + text[match.end():]
            print(f"Found and redacted Aadhaar number.")
        else:
            # This case should be rare with the refined regex, but good to keep
            print(f"Pattern matched '{match.group(1)}' but length is not 12 after removing spaces.")
    else:
        print("Aadhaar number pattern not found for redaction.")
    return found_number, redacted_text

def analyze_text_with_gemini(text, api_key, model_name, doc_type, placeholder):
    """Uses Google Gemini API to extract structured data from OCR text."""
    print(f"Analyzing text for document type '{doc_type}' using Gemini model '{model_name}'...")
    if genai is None:
        print("Error: google-generativeai library is not available.")
        return None
    if not api_key or api_key == "YOUR_GEMINI_API_KEY": # Check against default placeholder too
        print("Error: Gemini API Key is not configured.")
        return None

    try:
        # Configure the Gemini client
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(model_name)

        # Define prompts based on document type
        if doc_type == "passport":
            prompt = f"""Analyze the following text extracted from an INDIAN PASSPORT. The passport number has been redacted and replaced with '{placeholder}'.
Extract the following fields: Type, Country Code, Surname, Given Names, Nationality, Sex, Date of Birth, Place of Birth, Place of Issue, Date of Issue, Date of Expiry.
Format the output STRICTLY as a JSON object. Use null or an empty string "" if a field cannot be found. Respond ONLY with the JSON object.

Text:
---
{text}
---
JSON Output:"""
        elif doc_type == "aadhaar":
            prompt = f"""Analyze the following text extracted from an INDIAN AADHAAR CARD. The Aadhaar number has been redacted and replaced with '{placeholder}'.
Extract the following fields: Name, Date of Birth OR Year of Birth (prefer Date of Birth if available), Sex/Gender, Full Address (including Pincode if present), Date of Issue (if available).
Format the output STRICTLY as a JSON object. Use null or an empty string "" if a field cannot be found. Respond ONLY with the JSON object.

Text:
---
{text}
---
JSON Output:"""
        else:
            print(f"Error: Unsupported document type '{doc_type}' for Gemini analysis.")
            return None

        # Send the request to Gemini
        print("Sending request to Gemini API...")
        response = model.generate_content(prompt)

        # Process the response
        if not response.parts:
            print("Gemini Error: The response did not contain any parts.")
            # Log safety feedback if available
            if hasattr(response, 'prompt_feedback'): print(f"Prompt Feedback: {response.prompt_feedback}")
            return None

        print("Gemini analysis successful. Received response.")
        # Assuming the response text is the JSON string
        return response.text

    except Exception as e:
        print(f"Error during Gemini API call or processing: {e}")
        print(traceback.format_exc())
        # Check for specific API key related errors if possible (though the exception message might be generic)
        if "API key not valid" in str(e):
             print("Hint: The Gemini API key might be invalid or missing permissions.")
        return None

def compute_unique_identifiers(data, doc_type):
    """Computes a unique combination of name and DOB/YOB for user matching."""
    name, dob_yob = None, None
    print(f"Computing unique identifiers for doc_type: {doc_type}")
    try:
        if not isinstance(data, dict):
            print("Error: Input data is not a dictionary.")
            return None, None

        if doc_type == "passport":
            surname = data.get("Surname", "").strip()
            given_names = data.get("Given Names", "").strip()
            # Combine names, handling cases where one might be missing
            if given_names and surname:
                name = f"{given_names} {surname}".lower()
            elif given_names:
                name = given_names.lower()
            elif surname:
                name = surname.lower() # Less common, but possible
            # Get Date of Birth, checking common key variations
            dob_yob = (data.get("Date of Birth") or data.get("Date_of_Birth", "")).strip()

        elif doc_type == "aadhaar":
            name = data.get("Name", "").strip().lower()
            # Prioritize full Date of Birth over Year of Birth
            dob_raw = data.get("Date of Birth") or data.get("Date_of_Birth") or \
                       data.get("Year of Birth") or data.get("Year_of_Birth")
            # Handle potential nested structures if Gemini returns complex objects (unlikely with prompt)
            if isinstance(dob_raw, dict):
                 dob_raw = dob_raw.get("Date of Birth") or dob_raw.get("Date_of_Birth") or \
                           dob_raw.get("Year of Birth") or dob_raw.get("Year_of_Birth")
            dob_yob = (dob_raw or "").strip()

        if name and dob_yob:
            print(f"Computed identifiers - Name: '{name}', DOB/YOB: '{dob_yob}'")
            return name, dob_yob
        else:
            print(f"Warning: Could not compute sufficient identifiers (Name: '{name}', DOB/YOB: '{dob_yob}')")
            return None, None

    except Exception as e:
        print(f"Error computing unique identifiers: {e}")
        print(traceback.format_exc())
        return None, None

def extract_data_from_word(file_path):
    """Extracts Question-Answer pairs from the first table in a .docx file."""
    print(f"Extracting Q&A data from Word table: {file_path}")
    if docx is None:
        print("Error: python-docx library is not available.")
        return None
    qa_data = {}
    try:
        document = docx.Document(file_path)
        if not document.tables:
            print("Warning: No tables found in the Word document.")
            return {} # Return empty dict if no tables

        print(f"Processing first table out of {len(document.tables)} found.")
        table = document.tables[0]
        is_header_row = True # Assume first row is header
        row_index = 0
        for row in table.rows:
            row_index += 1
            if is_header_row:
                is_header_row = False # Skip the header row
                continue

            # Assuming Question is in the 2nd cell (index 1) and Answer in the 3rd (index 2)
            if len(row.cells) >= 3:
                question = row.cells[1].text.strip()
                answer = row.cells[2].text.strip()
                if question: # Only add if question is not empty
                    if question in qa_data:
                        # Handle duplicate questions if necessary (e.g., log warning, overwrite, append)
                        print(f"Warning: Duplicate question found at row {row_index}: '{question}'. Overwriting previous answer.")
                    qa_data[question] = answer
                else:
                     print(f"Warning: Row {row_index} has an empty question cell (cell index 1).")
            else:
                print(f"Warning: Row {row_index} has fewer than 3 cells ({len(row.cells)}). Skipping.")

        if not qa_data:
            print("Warning: No Question-Answer pairs were extracted from the table.")
        else:
            print(f"Successfully extracted {len(qa_data)} Q&A pairs.")
        return qa_data

    except Exception as e:
        print(f"Error extracting data from Word table in '{file_path}': {e}")
        print(traceback.format_exc())
        return None # Return None on critical error

def sanitize_mongodb_keys(obj):
    """Recursively sanitizes keys in a dictionary or list for MongoDB compatibility."""
    if isinstance(obj, dict):
        new_dict = {}
        for key, value in obj.items():
            # Ensure key is a string
            if not isinstance(key, str):
                try: key = str(key)
                except: continue # Skip if key cannot be converted to string

            if not key: continue # Skip empty keys

            # Replace '.' with '_' and remove leading '$'
            sanitized_key = key.replace('.', '_')
            if sanitized_key.startswith('$'):
                sanitized_key = '_' + sanitized_key[1:]

            if not sanitized_key: continue # Skip if key becomes empty after sanitization

            # Recursively sanitize value
            new_dict[sanitized_key] = sanitize_mongodb_keys(value)
        return new_dict
    elif isinstance(obj, list):
        # Recursively sanitize items in a list
        return [sanitize_mongodb_keys(item) for item in obj]
    else:
        # Return non-dict/list items as is
        return obj

def split_name(full_name):
    """Splits a full name string into surname and given names."""
    if not full_name or not isinstance(full_name, str):
        return "", ""
    parts = full_name.strip().split()
    if not parts:
        return "", ""
    if len(parts) == 1:
        # Assume single part is given name if only one part exists
        return "", parts[0].upper()
    else:
        # Assume last part is surname, rest is given names
        surname = parts[-1].upper()
        given_names = " ".join(parts[:-1]).upper()
        return surname, given_names

def format_date_for_form(date_str):
    """Parses various common date string formats and returns YYYY-MM-DD."""
    if not date_str or not isinstance(date_str, str):
        return None

    # List of expected date formats
    # Prioritize formats that are less ambiguous (e.g., with month names or 4-digit years)
    formats = [
        "%d %b %Y",  # 01 Jan 2023
        "%d-%b-%Y",  # 01-Jan-2023
        "%B %d, %Y", # January 01, 2023
        "%b %d, %Y",  # Jan 01, 2023
        "%Y-%m-%d",  # 2023-01-31 (ISO format)
        "%d/%m/%Y",  # 31/01/2023 (Common EU/India)
        "%m/%d/%Y",  # 01/31/2023 (Common US)
        "%Y",        # 2023 (Year only) - Default to Jan 1st
        # Add more formats if needed
         "%d%m%Y", # 01012023
         "%d %m %Y", # 01 01 2023
    ]

    cleaned_date_str = date_str.strip()
     # Remove ordinal suffixes (st, nd, rd, th) before parsing
    cleaned_date_str = re.sub(r"(\d+)(st|nd|rd|th)\b", r"\1", cleaned_date_str, flags=re.IGNORECASE)


    # Attempt to parse using the defined formats
    for fmt in formats:
        try:
            # Use strptime for parsing
            dt_object = datetime.strptime(cleaned_date_str, fmt)
            # Format the output as YYYY-MM-DD
            # If only year was parsed, default to YYYY-01-01
            return dt_object.strftime('%Y-%m-%d') if fmt != "%Y" else dt_object.strftime('%Y-01-01')
        except ValueError:
            continue # Try the next format if parsing fails

    # Fallback check for just a 4-digit year string
    if re.fullmatch(r"\d{4}", cleaned_date_str):
        return f"{cleaned_date_str}-01-01"

    # If no format matches
    print(f"Warning: Date parsing failed for input '{date_str}'. Could not match any known format.")
    return None # Return None if parsing fails for all formats


def format_date_ddmmyyyy(date_str):
    """Parses various common date string formats and returns DD-MM-YYYY."""
    if not date_str or not isinstance(date_str, str):
        return None

    # List of expected date formats (similar to format_date_for_form)
    formats = [
        "%d %b %Y", "%d-%b-%Y", "%B %d, %Y", "%b %d, %Y",
        "%Y-%m-%d", "%d/%m/%Y", "%m/%d/%Y", "%Y",
        "%d%b%Y", # e.g., 15APR2026
        "%d %B %Y", # e.g., 15 April 2026
        "%dth %B %Y", # e.g., 15th April 2026 - Handle 'th', 'st', 'nd', 'rd' (handled by removal below)
        "%dst %B %Y",
        "%dnd %B %Y",
        "%drd %B %Y",
         "%d%m%Y", # 01012023
         "%d %m %Y", # 01 01 2023
    ]

    cleaned_date_str = date_str.strip()

    # Remove ordinal suffixes (st, nd, rd, th) before parsing
    cleaned_date_str = re.sub(r"(\d+)(st|nd|rd|th)\b", r"\1", cleaned_date_str, flags=re.IGNORECASE)

    # Attempt to parse using the defined formats
    for fmt in formats:
        try:
            dt_object = datetime.strptime(cleaned_date_str, fmt)
            # Format the output as DD-MM-YYYY
            # If only year was parsed, default to 01-01-YYYY
            return dt_object.strftime('%d-%m-%Y') if fmt != "%Y" else dt_object.strftime('01-01-%Y')
        except ValueError:
            continue

    # Fallback check for just a 4-digit year string
    if re.fullmatch(r"\d{4}", cleaned_date_str):
        return f"01-01-{cleaned_date_str}"

    # If no format matches
    print(f"Warning: Date parsing failed for DD-MM-YYYY output for input '{date_str}'.")
    return None # Return None if parsing fails

def parse_place_of_birth(pob_str):
    """Parses a place of birth string into city and state/country."""
    if not pob_str or not isinstance(pob_str, str):
        return None, None
    # Split by comma or space, remove empty parts
    parts = [p.strip() for p in re.split(r'\s*,\s*|\s+', pob_str) if p.strip()]
    if len(parts) >= 2:
        # Assume first part is city, second is state/country
        return parts[0].upper(), parts[1].upper()
    elif len(parts) == 1:
        # Assume only city is provided
        return parts[0].upper(), None
    else:
        # No usable parts found
        return None, None


# --- Flask Routes ---

@app.route('/', methods=['GET'])
def index():
    """Renders the main upload page."""
    if mongo_client is None:
        # Provide a user-friendly error message if DB is down
        return "Error: Database connection is unavailable. Please try again later.", 503
    print("Rendering homepage.")
    # Create the allowed extensions string for the template
    allowed_ext_str = ", ".join(f".{ext}" for ext in sorted(list(ALLOWED_EXTENSIONS)))
    # Pass status and allowed extensions to the template
    return render_template('index.html',
                           allowed_ext_str=allowed_ext_str, # Pass the string here
                           db_status="Connected" if mongo_client else "Disconnected")


# --- Main Upload Endpoint ---
@app.route('/upload', methods=['POST'])
def upload():
    """Handles file uploads (Images and Word docs), processes them, and stores data."""
    print("\n--- New Upload Request Received ---")
    start_time = datetime.now()
    file_path = None # Define file_path here to ensure it's available in finally block

    try: # Wrap main logic in try...finally for cleanup
        # Check prerequisites
        if mongo_client is None or db is None:
            print("Error: MongoDB connection not available.")
            return jsonify({"error": "Database service is currently unavailable."}), 503
        if 'file' not in request.files:
            print("Error: No file part in the request.")
            return jsonify({"error": "No file part received."}), 400
        file = request.files['file']
        if file.filename == '':
            print("Error: No file selected.")
            return jsonify({"error": "No file selected for upload."}), 400

        # Secure filename and get extension
        safe_filename = secure_filename(file.filename)
        _, file_ext = os.path.splitext(safe_filename)
        file_ext = file_ext.lower()
        print(f"Processing uploaded file: {safe_filename} (Type: {file_ext})")

        # *** Get the document type selected by the user from the form data ***
        selected_type_from_frontend = request.form.get("selected_doc_type")
        if not selected_type_from_frontend:
            print("Error: 'selected_doc_type' missing from form data.")
            return jsonify({"error": "Internal configuration error: Missing document type selection from frontend."}), 400
        print(f"Selected document type from frontend: {selected_type_from_frontend}")

        # *** Basic validation: Check file extension against selected type ***
        if selected_type_from_frontend in ["passport", "aadhaar"] and file_ext not in ['.png', '.jpg', '.jpeg']:
             print(f"Validation Error: Frontend selected '{selected_type_from_frontend}' but file is not an image ({file_ext}).")
             return jsonify({"error": f"File Type Mismatch: Expected an image file (.png, .jpg, .jpeg) for selected type '{selected_type_from_frontend}', but received '{file_ext}'."}), 400
        if selected_type_from_frontend == "word_form" and file_ext != '.docx':
             print(f"Validation Error: Frontend selected 'word_form' but file is not '.docx' ({file_ext}).")
             return jsonify({"error": f"File Type Mismatch: Expected a Word document (.docx) for selected type 'Word Form Questionnaire', but received '{file_ext}'."}), 400
        # Note: The allowed_file check below is still useful as a general safeguard
        if not allowed_file(safe_filename):
            print(f"Error: File type not allowed by general filter: {file.filename}")
            return jsonify({"error": f"Invalid file format. Allowed formats: {ALLOWED_EXTENSIONS}"}), 415


        # Initialize variables
        # file_path defined outside try block now
        final_response_data = {} # Holds data for the JSON response
        user_id = None # Will store the ObjectId of the user
        doc_type = "unknown" # Backend detected type (passport, aadhaar, word_form)
        extracted_data_dict = {} # Holds the structured data extracted from the file
        current_scan_results = {} # Temporary storage for intermediate results (like found IDs)

        # Save the uploaded file temporarily
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], safe_filename)
        file.save(file_path)
        print(f"File saved temporarily to: {file_path}")

        # --- IMAGE PROCESSING BRANCH ---
        if file_ext in ['.png', '.jpg', '.jpeg']:
            print("Processing as Image file...")
            # Ensure frontend selection matches image processing branch
            if selected_type_from_frontend not in ["passport", "aadhaar"]:
                 print(f"Validation Error: Image file uploaded, but frontend selected type was '{selected_type_from_frontend}'.")
                 return jsonify({"error": f"File Type Mismatch: Received an image file, but expected type was '{selected_type_from_frontend}' based on selection."}), 400

            id_field_name = "Document_ID_Number" # Default key for the ID number

            # 1. Perform OCR
            if PaddleOCR is None:
                print("Error: OCR dependency (PaddleOCR) not available.")
                return jsonify({"error": "Image processing engine is not available."}), 501 # Not Implemented
            full_ocr_text = extract_text_paddle(file_path, lang='en') # Assuming English
            if full_ocr_text is None: # Critical OCR failure
                print("Error: OCR process failed critically.")
                return jsonify({"error": "Failed to extract text from image."}), 500
            if not full_ocr_text.strip(): # OCR ran but found no text
                print("Warning: OCR completed but found no text in the image.")
                return jsonify({"error": "Could not find any text in the uploaded image."}), 400

            # 2. Detect Document Type (Backend Detection)
            doc_type = detect_document_type(full_ocr_text) # This is backend's detected type

            # *** VALIDATION STEP: Compare frontend selection with backend detection ***
            print(f"Backend detected type: '{doc_type}'. Validating against frontend selection: '{selected_type_from_frontend}'...")

            if selected_type_from_frontend == "passport" and doc_type != "passport":
                print(f"Validation Error: User selected Passport, but backend detected '{doc_type}'.")
                return jsonify({"error": f"Content Mismatch: The uploaded image was not recognized as a Passport, although 'Passport Image' was selected. Detected content seems to be: '{doc_type.capitalize()}'."}), 400

            if selected_type_from_frontend == "aadhaar" and doc_type != "aadhaar":
                print(f"Validation Error: User selected Aadhaar, but backend detected '{doc_type}'.")
                return jsonify({"error": f"Content Mismatch: The uploaded image was not recognized as an Aadhaar card, although 'Aadhaar Image' was selected. Detected content seems to be: '{doc_type.capitalize()}'."}), 400

            # Also handle if backend detection failed altogether (doc_type is 'unknown')
            if doc_type == "unknown":
                 print(f"Validation Error: Backend could not determine document type ('{doc_type}') for selected type '{selected_type_from_frontend}'.")
                 ocr_preview = full_ocr_text[:500] + '...' if len(full_ocr_text) > 500 else full_ocr_text
                 return jsonify({
                     "error": f"Content Mismatch: Could not identify the document type from the image content for the selected type '{selected_type_from_frontend.capitalize()}'. Please ensure the image is clear.",
                     "ocr_text_preview": ocr_preview
                 }), 400
            # --- End Validation Step ---
            print("Validation successful: Frontend selection matches backend detection.")

            # If validation passes, proceed...
            final_response_data["Detected Document Type"] = doc_type.capitalize()
            id_field_name = f"{doc_type.capitalize()}_No" # e.g., Passport_No, Aadhaar_No

            # 3. Find and Redact ID Number (Based on backend detected type 'doc_type')
            redacted_text = full_ocr_text
            found_id = None
            placeholder = "[REDACTED ID]"
            if doc_type == "passport":
                found_id, redacted_text = find_and_redact_passport_no(full_ocr_text)
                placeholder = PASSPORT_REDACTION_PLACEHOLDER
            elif doc_type == "aadhaar":
                found_id, redacted_text = find_and_redact_aadhaar_no(full_ocr_text)
                placeholder = AADHAAR_REDACTION_PLACEHOLDER
            current_scan_results[id_field_name] = found_id if found_id else ID_NUMBER_NOT_FOUND

            # 4. Analyze with Gemini (using backend detected type 'doc_type')
            gemini_output = None
            gemini_status = "Skipped (Not Available/Configured)"
            gemini_parsed_data = {}
            if genai and API_KEY and API_KEY != "YOUR_GEMINI_API_KEY": # Added check for empty key
                gemini_output = analyze_text_with_gemini(redacted_text, API_KEY, GEMINI_MODEL_NAME, doc_type, placeholder)
                if gemini_output:
                    # Clean the response to extract JSON
                    clean_response = gemini_output.strip().removeprefix("```json").removeprefix("```").removesuffix("```").strip()
                    json_start = clean_response.find('{')
                    json_end = clean_response.rfind('}')
                    if json_start != -1 and json_end != -1 and json_start < json_end:
                        json_string = clean_response[json_start : json_end + 1]
                        try:
                            parsed = json.loads(json_string)
                            if isinstance(parsed, dict):
                                gemini_parsed_data = parsed
                                gemini_status = "Success"
                                print("Successfully parsed Gemini JSON response.")
                            else:
                                raise TypeError("Parsed JSON is not a dictionary.")
                        except (json.JSONDecodeError, TypeError) as e:
                            print(f"Error parsing Gemini JSON response: {e}")
                            print(f"Received response snippet: {clean_response[:200]}...")
                            gemini_status = f"Failed: JSON Parse Error ({e})"
                    else:
                        print("Error: Could not find valid JSON object in Gemini response.")
                        print(f"Received response snippet: {clean_response[:200]}...")
                        gemini_status = "Failed: No JSON object found in response"
                else:
                    gemini_status = "Failed: No response or error during analysis"
            final_response_data["gemini_analysis_status"] = gemini_status

            # 5. Compute Unique Identifiers (using backend detected type 'doc_type')
            computed_name, computed_dob_yob = compute_unique_identifiers(gemini_parsed_data, doc_type)

            # 6. Find or Create User Record in MongoDB
            # (User creation/lookup logic remains the same)
            final_response_data["db_user_status"] = "Skipped"
            if users_collection is not None and computed_name and computed_dob_yob:
                try:
                    print(f"Attempting to find or create user: Name='{computed_name}', DOB/YOB='{computed_dob_yob}'")
                    user_filter = {"name": computed_name, "dob_yob": computed_dob_yob}
                    user_update = {
                        "$setOnInsert": {
                            "name": computed_name,
                            "dob_yob": computed_dob_yob,
                            "created_at": datetime.utcnow()
                        }
                    }
                    user_record = users_collection.find_one_and_update(
                        user_filter,
                        user_update,
                        upsert=True,
                        return_document=ReturnDocument.AFTER
                    )
                    if user_record and '_id' in user_record:
                        user_id = user_record['_id']
                        print(f"User found or created successfully. User ID: {user_id}")
                        final_response_data["db_user_status"] = "Success (Found/Created via Image Data)"
                        final_response_data["user_id"] = str(user_id)
                    else:
                        print("Error: User upsert operation did not return a valid record.")
                        final_response_data["db_user_status"] = "Failed (DB Upsert Error)"
                except pymongo.errors.DuplicateKeyError:
                     print("Warning: Potential race condition during user upsert. Attempting to fetch existing user.")
                     user_record = users_collection.find_one(user_filter)
                     if user_record:
                         user_id = user_record['_id']
                         final_response_data["db_user_status"] = "Success (Fetched Existing User after Race Condition)"
                         final_response_data["user_id"] = str(user_id)
                     else:
                         print("Error: Could not fetch user even after duplicate key error.")
                         final_response_data["db_user_status"] = "Failed (Duplicate Key Resolution Error)"
                except Exception as e:
                    print(f"Error during user find/create operation: {e}")
                    print(traceback.format_exc())
                    final_response_data["db_user_status"] = "Failed (Database Error)"
            elif not computed_name or not computed_dob_yob:
                final_response_data["db_user_status"] = "Skipped (Insufficient Identifiers from Image)"
                print("Skipping user find/create: Insufficient name/DOB info extracted.")
            elif users_collection is None:
                 final_response_data["db_user_status"] = "Skipped (DB Not Connected)"


            # Prepare the final dictionary of extracted data for this image
            extracted_data_dict = gemini_parsed_data.copy()
            if found_id and id_field_name:
                 sanitized_id_key = id_field_name.replace('.', '_').replace(' ', '_')
                 extracted_data_dict[sanitized_id_key] = found_id
                 print(f"Added found ID to extracted data: {sanitized_id_key}={found_id}")


        # --- WORD DOCUMENT PROCESSING BRANCH ---
        elif file_ext == '.docx':
            print("Processing as Word Document (.docx)...")
            # Ensure frontend selection matches word processing branch
            if selected_type_from_frontend != "word_form":
                 print(f"Validation Error: Word file uploaded, but frontend selected type was '{selected_type_from_frontend}'.")
                 return jsonify({"error": f"File Type Mismatch: Received a Word file (.docx), but expected type was '{selected_type_from_frontend}' based on selection."}), 400

            # Use the backend constant WORD_DOC_TYPE for internal logic
            doc_type = WORD_DOC_TYPE
            final_response_data["Detected Document Type"] = "Word Form" # Keep user-friendly name for response
            final_response_data["gemini_analysis_status"] = "N/A"

            # Check if docx library is available
            if docx is None:
                print("Error: python-docx library is missing. Cannot process Word files.")
                return jsonify({"error": "Backend configuration error: Cannot process Word documents."}), 500

            # 1. Extract Full Text (primarily for finding Passport No)
            word_full_text = extract_full_text_from_word(file_path)
            if word_full_text is None:
                print("Error: Failed to extract any text from the Word document.")
                return jsonify({"error": "Failed to read text content from the Word document."}), 500
            if not word_full_text:
                 print("Warning: Word document appears to be empty or contains no text.")
                 return jsonify({"error": "The uploaded Word document is empty or contains no text."}), 400

            # 2. Find Passport Number within the Word document text
            passport_no_from_word = find_passport_no_in_text(word_full_text)
            if not passport_no_from_word:
                print(f"Error: Passport number could not be found within the text of '{safe_filename}'.")
                return jsonify({"error": f"Passport number pattern not found in the uploaded Word document '{safe_filename}'. Cannot link document to a user."}), 400 # Cannot link without passport#
            final_response_data["passport_no_found_in_word"] = passport_no_from_word
            current_scan_results["passport_number_from_text"] = passport_no_from_word

            # 3. Look up User ID based on the found Passport Number
            # (User lookup logic remains the same)
            final_response_data["db_user_status"] = "Pending Lookup"
            if documents_collection is None:
                 return jsonify({"error": "Database service is currently unavailable."}), 503
            try:
                print(f"Searching for existing Passport document with number: {passport_no_from_word} to find user ID...")
                # Query the documents collection for a passport type doc with matching number
                lookup_filter = {"doc_type": "passport", "extracted_data.Passport_No": passport_no_from_word}
                matched_passport_doc = documents_collection.find_one(lookup_filter)

                if matched_passport_doc:
                    linked_user_id = matched_passport_doc.get("user_id")
                    if linked_user_id and isinstance(linked_user_id, ObjectId):
                        user_id = linked_user_id
                        print(f"User ID found via linked Passport document: {user_id}")
                        final_response_data["db_user_status"] = "Success (Found via Passport No Lookup)"
                        final_response_data["user_id"] = str(user_id)
                    else:
                        print(f"Error: Found passport doc (ID: {matched_passport_doc.get('_id')}) but it has missing or invalid user_id link.")
                        return jsonify({"error": "Database integrity issue: Found matching passport record but cannot link to a user."}), 500
                else:
                    print(f"Error: No existing Passport record found for number '{passport_no_from_word}'. Cannot link Word form to a user.")
                    return jsonify({"error": f"User lookup failed: No passport record found matching the number '{passport_no_from_word}' in the Word document. Please upload the corresponding passport image first."}), 404 # Not Found

            except Exception as e:
                print(f"Database error during user lookup via passport number: {e}")
                print(traceback.format_exc())
                return jsonify({"error": "Database error during user lookup."}), 500

            # 4. Extract Question-Answer Data from the Word table
            if not user_id:
                 print("Error: User ID missing after lookup attempt. Cannot proceed.")
                 return jsonify({"error": "Internal error: User ID not established."}), 500
            raw_qa_data = extract_data_from_word(file_path)
            if raw_qa_data is None:
                 print("Error: Failed to extract Q&A data from the Word document table.")
                 return jsonify({"error": "Failed to extract question-answer data from the Word document."}), 500
            if not raw_qa_data:
                 print("Warning: No Q&A pairs extracted from the Word document table.")

            # 5. Map Extracted Questions to Standardized Keys
            # (Mapping and secondary parsing logic remains the same)
            mapped_data = {}
            skipped_questions = []
            print("Mapping extracted questions to standardized keys...")
            for question, answer in raw_qa_data.items():
                mapped_key = QUESTION_TO_KEY_MAP.get(question)
                if mapped_key:
                    mapped_data[mapped_key] = answer
                else:
                    skipped_questions.append(question)
                    print(f"  -> Skipping unmapped question: '{question[:100]}...'") # Log snippet
            if skipped_questions:
                final_response_data["warning"] = f"Skipped {len(skipped_questions)} unmapped questions found in the Word document."

            print("Performing secondary parsing for specific fields (payer, parents)...")
            data_to_save = mapped_data.copy() # Work on a copy

            # --- >>> FORMAT US DATES TO DD-MM-YYYY <<< ---
            print("Formatting US arrival/departure dates to DD-MM-YYYY for storage...")
            arrival_key = 'us_arrival_date'
            departure_key = 'us_departure_date'

            if arrival_key in data_to_save and data_to_save[arrival_key]:
                original_arrival = data_to_save[arrival_key]
                formatted_arrival = format_date_ddmmyyyy(original_arrival)
                if formatted_arrival:
                    data_to_save[arrival_key] = formatted_arrival
                    print(f"  -> Formatted '{arrival_key}': '{original_arrival}' -> '{formatted_arrival}'")
                else:
                    print(f"  -> Warning: Could not format '{arrival_key}' ('{original_arrival}') to DD-MM-YYYY. Storing original value.")

            if departure_key in data_to_save and data_to_save[departure_key]:
                original_departure = data_to_save[departure_key]
                formatted_departure = format_date_ddmmyyyy(original_departure)
                if formatted_departure:
                    data_to_save[departure_key] = formatted_departure
                    print(f"  -> Formatted '{departure_key}': '{original_departure}' -> '{formatted_departure}'")
                else:
                    print(f"  -> Warning: Could not format '{departure_key}' ('{original_departure}') to DD-MM-YYYY. Storing original value.")
            # --- >>> END DATE FORMATTING <<< ---

            if passport_no_from_word and 'passport_number' not in data_to_save:
                print(f"Adding passport number ('{passport_no_from_word}') found in text to final data.")
                data_to_save['passport_number'] = passport_no_from_word
            elif 'passport_number' in data_to_save and data_to_save['passport_number'] != passport_no_from_word:
                 print(f"Warning: Passport number in Q&A ('{data_to_save['passport_number']}') differs from number found in text ('{passport_no_from_word}'). Using Q&A value.")


            extracted_data_dict = data_to_save
            if not extracted_data_dict:
                print("Warning: No data mapped or added after processing Word document.")

        else:
            # This case should be caught by earlier checks, but acts as a fallback
            print(f"Error: Unsupported file extension encountered: '{file_ext}' after initial checks.")
            return jsonify({"error": f"Unsupported file type '{file_ext}'. Allowed types: {ALLOWED_EXTENSIONS}"}), 415


        # --- COMMON STEPS (executed after either Image or Word processing, if validation passed) ---

        # 6. Sanitize Final Data Dictionary Keys for MongoDB Compatibility
        # (Sanitization logic remains the same)
        sanitized_data_to_save = {}
        try:
            if not isinstance(extracted_data_dict, dict):
                print("Warning: Extracted data is not a dictionary. Resetting to empty dict.")
                extracted_data_dict = {}
            print("Sanitizing final data keys for MongoDB storage...")
            sanitized_data_to_save = sanitize_mongodb_keys(extracted_data_dict)
            print("Key sanitization complete.")
            if not sanitized_data_to_save and extracted_data_dict:
                print("Warning: Data dictionary became empty after key sanitization.")
        except Exception as e:
            print(f"Error during data key sanitization: {e}")
            print(traceback.format_exc())
            return jsonify({"error": "Internal server error during data preparation."}), 500


        # 7. Store Document Data in MongoDB (if user_id is available)
        # (DB storage logic remains the same, using backend detected `doc_type`)
        final_response_data["db_document_status"] = "Pending"
        if user_id and documents_collection is not None:
            try:
                print(f"Checking for existing document of type '{doc_type}' for user: {user_id}")
                # Check if a document of the same type already exists for this user
                existing_doc = documents_collection.find_one({"user_id": user_id, "doc_type": doc_type})

                if existing_doc:
                    existing_doc_id = str(existing_doc.get('_id'))
                    print(f"Document type '{doc_type}' already exists for this user (ID: {existing_doc_id}). Skipping insertion.")
                    final_response_data["db_document_status"] = "Skipped (Already Exists)"
                    final_response_data["message"] = f"A '{doc_type.replace('_', ' ').capitalize()}' document record already exists for this user."
                    final_response_data["existing_document_record_id"] = existing_doc_id
                    final_response_data["extracted_data_preview"] = existing_doc.get("extracted_data", {})

                else:
                    document_record = {
                        "user_id": user_id, # Store as ObjectId
                        "doc_type": doc_type, # Store the backend *detected* type
                        "original_filename": safe_filename,
                        "scan_timestamp_utc": datetime.utcnow(),
                        "extracted_data": sanitized_data_to_save if isinstance(sanitized_data_to_save, dict) else {},
                        "gemini_analysis_status": final_response_data.get("gemini_analysis_status", "N/A")
                    }
                    print("Inserting new document record into MongoDB...")
                    insert_result = documents_collection.insert_one(document_record)
                    inserted_doc_id = insert_result.inserted_id
                    print(f"Document record inserted successfully (ID: {inserted_doc_id})")
                    final_response_data["db_document_status"] = "Success"
                    final_response_data["document_record_id"] = str(inserted_doc_id)
                    final_response_data["message"] = f"Successfully added '{doc_type.replace('_', ' ').capitalize()}' document data."
                    if sanitized_data_to_save:
                         final_response_data["extracted_data_preview"] = sanitized_data_to_save

            except Exception as e:
                print(f"Database error during document check or insertion: {e}")
                print(traceback.format_exc())
                final_response_data["db_document_status"] = "Failed (Database Error)"
                final_response_data["message"] = "An error occurred while saving the document data."

        elif not user_id:
            final_response_data["db_document_status"] = "Skipped (User ID Missing or Not Found)"
            print("Skipping document storage: User ID was not established.")
            if not final_response_data.get("message"):
                 final_response_data["message"] = "Could not store document data because the user could not be identified or linked."
        elif documents_collection is None:
             final_response_data["db_document_status"] = "Skipped (DB Error)"


        # --- Prepare Final JSON Response ---
        # (Response preparation logic remains the same)
        end_time = datetime.now()
        processing_time = (end_time - start_time).total_seconds()
        final_response_data["_processing_time_seconds"] = round(processing_time, 2)

        cleaned_response = {}
        response_keys = [
            "user_id", "document_record_id", "existing_document_record_id",
            "db_user_status", "db_document_status", "message",
            "_processing_time_seconds", "warning", "passport_no_found_in_word",
            "Detected Document Type", "extracted_data_preview", "gemini_analysis_status"
        ]
        if 'id_field_name' in locals() and id_field_name and id_field_name in current_scan_results:
             response_keys.append(id_field_name)
             final_response_data[id_field_name] = current_scan_results[id_field_name]

        for key in response_keys:
            if key in final_response_data and final_response_data[key] is not None:
                if key == "extracted_data_preview" and not final_response_data[key]:
                    continue
                cleaned_response[key] = final_response_data[key]

        print(f"Request processing finished in {processing_time:.2f} seconds.")
        status_code = 200 # Default to OK
        if final_response_data.get("db_document_status") == "Success":
            status_code = 201 # Use 201 Created for new resource
        elif final_response_data.get("db_document_status") == "Skipped (Already Exists)":
             status_code = 200 # OK is appropriate here

        return jsonify(cleaned_response), status_code

    except Exception as e:
        # Catch-all for any unexpected errors during the main processing block
        print(f"FATAL error occurred during /upload processing: {e}")
        print(traceback.format_exc())
        # Return a generic 500 error
        return jsonify({"error": "An unexpected internal server error occurred."}), 500
    finally:
        # Cleanup: Remove the temporarily saved file regardless of success or failure
        if file_path and os.path.exists(file_path):
            try:
                os.remove(file_path)
                print(f"Cleaned up temporary file: {file_path}")
            except OSError as e:
                # Log error but don't prevent response from being sent
                print(f"Error removing temporary file '{file_path}': {e}")


# --- API Endpoint to GET ALL USERS ---
# (get_users endpoint remains the same)
@app.route('/api/getUsers', methods=['GET'])
def get_users():
    """Retrieves a list of all users from the database."""
    print("\n--- API Request Received: /api/getUsers ---")
    if users_collection is None:
        print("Error: User collection (DB) not available.")
        return jsonify({"error": "Database service is currently offline."}), 503
    try:
        print("Fetching users from database...")
        # Fetch only necessary fields (_id, name, dob_yob) and sort by name
        users_cursor = users_collection.find({}, {"_id": 1, "name": 1, "dob_yob": 1}).sort("name", pymongo.ASCENDING)
        users_list = list(users_cursor) # Convert cursor to list
        print(f"Found {len(users_list)} users.")

        # Format the user list for the response
        formatted_user_list = []
        for user in users_list:
            user_id_str = str(user['_id'])
            name = user.get('name', 'N/A').title() # Default to N/A and title case
            dob_yob = user.get('dob_yob', '') # Default to empty string
            # Create a display name combining name and DOB/YOB if available
            display_name = f"{name} ({dob_yob})" if dob_yob else name
            formatted_user_list.append({"id": user_id_str, "display_name": display_name})

        return jsonify(formatted_user_list)
    except Exception as e:
        print(f"Error fetching users from database: {e}")
        print(traceback.format_exc())
        return jsonify({"error": "An error occurred while fetching user data."}), 500


# --- API Endpoint for Form Filling Data ---
# (get_form_data endpoint remains the same)
@app.route('/api/getFormData', methods=['GET'])
def get_form_data():
    """
    Retrieves and combines data from Word Form, Passport, Aadhaar, and User records
    for a specific user ID to pre-fill a form.
    Prioritizes Word Form data, then Passport, then Aadhaar, then User record defaults.
    """
    print("\n--- API Request Received: /api/getFormData ---")
    if mongo_client is None: # Check primary client connection
        print("Error: MongoDB client not connected.")
        return jsonify({"message": "Database service is currently offline."}), 503

    # Get user ID from query parameter
    user_id_str = request.args.get('id')
    if not user_id_str:
        print("Error: 'id' query parameter is missing.")
        return jsonify({"message": "Missing required 'id' parameter in the request."}), 400

    # Validate and convert user ID string to ObjectId
    try:
        user_id_obj = ObjectId(user_id_str)
    except Exception: # Catches InvalidId and other potential errors
        print(f"Error: Invalid user ID format received: '{user_id_str}'")
        return jsonify({"message": "Invalid user ID format provided."}), 400

    print(f"Requesting combined form data for User ID: {user_id_obj}")

    try:
        # 1. Fetch User Record
        if users_collection is None: return jsonify({"message": "User collection unavailable."}), 503
        user_data = users_collection.find_one({"_id": user_id_obj})
        if not user_data:
            print(f"Error: User with ID '{user_id_obj}' not found.")
            return jsonify({"message": f"User with ID {user_id_str} not found."}), 404 # Not Found

        # 2. Fetch Associated Documents (Passport, Aadhaar, Word Form)
        if documents_collection is None: return jsonify({"message": "Documents collection unavailable."}), 503
        docs_cursor = documents_collection.find({
            "user_id": user_id_obj,
            "doc_type": {"$in": ["passport", "aadhaar", WORD_DOC_TYPE]}
        })
        # Organize documents by type for easier access
        docs_by_type = {doc['doc_type']: doc.get("extracted_data", {}) for doc in docs_cursor}
        print(f"Found associated documents for user: {list(docs_by_type.keys())}")

        # Extract data from each source, defaulting to empty dict if not found
        passport_data = docs_by_type.get("passport", {})
        aadhaar_data = docs_by_type.get("aadhaar", {})
        word_form_data = docs_by_type.get(WORD_DOC_TYPE, {})

        # 3. Initialize Combined Data Structure
        # Define all possible keys expected by the form-filling target (e.g., Chrome extension)
        all_form_keys = [
            # Personal Info
            "surnames", "given_names", "full_name_native", "sex", "marital_status", "date_of_birth",
            "pob_city", "pob_state_province", "pob_country", "national_id",
            # Contact Info
            "home_address_line1","home_address_line2","home_address_city","home_address_state","home_address_zip", "mailing_address_same_as_home", "primary_phone", "secondary_phone", "email",
            # Passport & ID
            "passport_number", "Type", "Country_Code", "Nationality", "Place_of_Issue", "Date_of_Issue", "Date_of_Expiry",
            "Aadhaar_No",
            # Family Info
            "father_surname","father_given_name" ,"father_dob", "father_in_us", "mother_surname","mother_given_name" ,"mother_dob", "mother_in_us",
            "spouse_surname","spouse_given_name", "spouse_dob", "spouse_pob", "spouse_address",
            "immediate_relatives_in_us", "other_relatives_in_us",
            # Employment Info (Current)
            "current_employer_name","current_employer_duties", "present_employer_address_line2", "present_employer_address_line1","present_employer_address_city","present_employer_address_state","present_employer_address_zip_code","present_employer_phone_number","prevoius_employer_duties", "current_employment_start_date",
            # Previous Employment Info
            "previously_employed","prev_emp_1_name", "prev_emp_1_address_line1","prev_emp_1_address_line2","prev_emp_1_address_city","prev_emp_1_address_state","prev_emp_1_address_zip_code","prev_emp_1_address_country","prev_emp_1_address_phone", "prev_emp_1_job_title","prev_emp_1_supervisor_surname","prev_emp_1_supervisor_given_name","prev_emp_1_start_date","prev_emp_1_phone","prev_emp_1_end_date","prev_emp_1_duties",
            "prev_emp_2_name", "prev_emp_2_address", "prev_emp_2_phone", "prev_emp_2_job_title","prev_emp_2_supervisor_name","prev_emp_2_start_date","prev_emp_2_end_date","prev_emp_2_duties",
            # Education Info
            "attended_secondary_education_or_above","highest_qualification", "institution_name", "institution_address_line1","institution_address_line2","institution_address_city","institution_address_state","institution_address_zip_code","institution_address_country","education_institute_start_date","education_institute_end_date",
            # US Travel Info
            "us_arrival_date", "us_departure_date", "us_arrival_departure_cities", "us_flight_details",
            "us_visit_locations","us_stay_address_zip_code", "us_stay_address_line1","us_stay_address_line2","us_stay_address_city","us_stay_address_state",
            # Trip Funding
            "trip_payer_surname","trip_payer","trip_payer_given_name","trip_payer_relationship","trip_payer_telephone","trip_payer_email", "trip_payer_address_same_as_home", # Added relationship
            "payer_address_same_as_applicant","payer_address_street_line1","payer_address_street_line2","payer_address_city","payer_address_state","payer_address_zip_code","payer_address_country",
        
            # Travel Companions & History
            "travel_companions", "previous_us_visits", "previous_us_visa_issued", "previous_us_visa_refusal","previous_us_visa_refusal_details",
            "immigrant_petition_filed", "lost_stolen_passport_details",
            # US Contact
            "us_contact_name_or_org", "us_contact_surname","us_contact_given_name", "us_contact_relationship",
            "us_contact_street_address_line1","us_contact_street_address_line2","us_contact_city","us_contact_state","us_contact_zip", "us_contact_phone","us_contact_email", # Assuming phone/email are combined in Word form parsing
            # Miscellaneous
            "social_media_presence", "clan_or_tribe", "other_permanent_residency", "other_nationalities",
            "languages_spoken", "countries_visited_last_5_years", "organization_contributions", "monthly_income",
            # Form-specific flags/defaults
            "full_name_native_na", "other_names_used", "has_telecode", "pob_state_province_na"
        ]
        # Initialize all keys to None
        combined_data = {key: None for key in all_form_keys}
        # Set default values for flags/checkboxes if needed
        combined_data.update({
            "full_name_native_na": False,
            "other_names_used": "N", # Assuming 'N' for No as default
            "has_telecode": "N",    # Assuming 'N' for No as default
            "pob_state_province_na": False
        })

        # 4. Data Mapping Logic (Layering data sources)
        print("Mapping data sources (Priority: Word -> Passport -> Aadhaar -> User Record)...")

        # Layer 1: Populate from Word Form Data (Highest priority for fields it contains)
        for key in combined_data:
            if key in word_form_data and word_form_data[key] not in [None, ""]:
                combined_data[key] = word_form_data[key]

        # Layer 2: Populate from Passport Data (Fill missing/empty OR override specific fields like names)
        passport_key_variants = {
            "surnames": ("Surname",),
            "given_names": ("Given Names", "Given_Names"),
            "sex": ("Sex",),
            "date_of_birth": ("Date of Birth", "Date_of_Birth"),
            "pob_raw": ("Place of Birth", "Place_of_Birth"), # Temporary key for PoB parsing
            "Nationality": ("Nationality",),
            "Country_Code": ("Country Code", "Country_Code"),
            "passport_number": ("Passport_No", "passport_number"), # Check specific stored key
            "Type": ("Type",),
            "Place_of_Issue": ("Place of Issue", "Place_of_Issue"),
            "Date_of_Issue": ("Date of Issue", "Date_of_Issue"),
            "Date_of_Expiry": ("Date of Expiry", "Date_of_Expiry")
        }
        for form_key, passport_options in passport_key_variants.items():
            passport_value = None
            for p_key in passport_options:
                if p_key in passport_data and passport_data[p_key] not in [None, ""]:
                    passport_value = passport_data[p_key]
                    break

            if passport_value:
                should_overwrite = (combined_data.get(form_key) in [None, ""]) or \
                                   (form_key in ["surnames", "given_names"])

                if should_overwrite:
                    if form_key in ["surnames", "given_names"]:
                        combined_data[form_key] = str(passport_value).upper()
                    else:
                        combined_data[form_key] = passport_value

        if combined_data.get("pob_city") is None and combined_data.get("pob_raw"):
            city, state_country = parse_place_of_birth(combined_data["pob_raw"])
            if city:
                combined_data["pob_city"] = city
                if state_country and combined_data.get("pob_state_province") is None:
                    combined_data["pob_state_province"] = state_country
                    combined_data["pob_state_province_na"] = False
                elif not state_country and combined_data.get("pob_state_province") is None:
                     combined_data["pob_state_province_na"] = True

        if "pob_raw" in combined_data: del combined_data["pob_raw"]


        # Layer 3: Populate from Aadhaar Data (Fill missing/empty fields)
        aadhaar_key_variants = {
            "full_name_aadhaar": ("Name",), # Temporary key for name splitting if needed
            "sex": ("Sex/Gender", "Sex_Gender"), # Check variations
            "date_of_birth": ("Date of Birth", "Date_of_Birth", "Year of Birth", "Year_of_Birth"), # DOB/YOB variations
            "Aadhaar_No": ("Aadhaar_No",), # Assuming stored with underscore
            "home_address": ("Address", "Full Address", "Full_Address") # Address variations
        }

        aadhaar_dob_value = None
        for dob_key in ("Date of Birth", "Date_of_Birth"):
            if dob_key in aadhaar_data and aadhaar_data[dob_key] not in [None, ""]:
                aadhaar_dob_value = aadhaar_data[dob_key]
                break
        if not aadhaar_dob_value:
            for yob_key in ("Year of Birth", "Year_of_Birth"):
                 if yob_key in aadhaar_data and aadhaar_data[yob_key] not in [None, ""]:
                     aadhaar_dob_value = aadhaar_data[yob_key]
                     break
        if combined_data.get("date_of_birth") in [None, ""] and aadhaar_dob_value:
            combined_data["date_of_birth"] = aadhaar_dob_value

        for form_key, aadhaar_options in aadhaar_key_variants.items():
            if form_key == "date_of_birth": continue

            aadhaar_value = None
            for a_key in aadhaar_options:
                if a_key in aadhaar_data and aadhaar_data[a_key] not in [None, ""]:
                    aadhaar_value = aadhaar_data[a_key]
                    break

            if combined_data.get(form_key) in [None, ""] and aadhaar_value:
                 combined_data[form_key] = aadhaar_value

        if combined_data.get("surnames") is None and combined_data.get("full_name_aadhaar"):
             print("  -> Using Aadhaar Name for splitting into surname/given names.")
             s_name, g_names = split_name(combined_data["full_name_aadhaar"])
             if s_name or g_names:
                 combined_data["surnames"] = s_name
                 combined_data["given_names"] = g_names
        if "full_name_aadhaar" in combined_data: del combined_data["full_name_aadhaar"]


        # Layer 4: Populate from User Record (Fill missing/empty as last resort)
        user_mapping = {
            "name": "user_record_name", # Temporary key for name splitting if needed
            "dob_yob": "date_of_birth", # Use user record DOB if still missing
            "full_name_native": "full_name_native",
            "full_name_native_na": "full_name_native_na",
            "other_names_used": "other_names_used",
            "has_telecode": "has_telecode",
            "pob_city": "pob_city",
            "pob_state_province": "pob_state_province",
            "pob_state_province_na": "pob_state_province_na",
            "pob_country": "pob_country"
        }
        for user_key, form_key in user_mapping.items():
            user_value = user_data.get(user_key)
            if user_value is not None:
                current_form_value = combined_data.get(form_key)

                use_user_value = False
                if isinstance(current_form_value, bool) and current_form_value is False:
                    if user_value: use_user_value = True
                elif current_form_value in [None, "", "N"]:
                    use_user_value = True

                if use_user_value:
                    if form_key in ["other_names_used", "has_telecode"] and isinstance(user_value, str):
                         combined_data[form_key] = user_value.upper()
                    else:
                         combined_data[form_key] = user_value

        if combined_data.get("surnames") is None and combined_data.get("user_record_name"):
             print("  -> Using User record Name for splitting.")
             s_name, g_names = split_name(combined_data["user_record_name"])
             if s_name or g_names:
                 combined_data["surnames"] = s_name
                 combined_data["given_names"] = g_names
        if "user_record_name" in combined_data: del combined_data["user_record_name"]


        # 5. Final Formatting and Cleanup
        print("Applying final formatting (dates, flags)...")
        date_keys_to_format_yyyy_mm_dd = [
            "date_of_birth", "spouse_dob", "father_dob", "mother_dob",
            "Date_of_Issue", "Date_of_Expiry"
            # DO NOT include us_arrival_date, us_departure_date here as they are DD-MM-YYYY
        ]
        for date_key in date_keys_to_format_yyyy_mm_dd:
            original_date = combined_data.get(date_key)
            if original_date:
                formatted_date = format_date_for_form(str(original_date))
                if formatted_date and formatted_date != original_date:
                     print(f"  -> Formatting date '{date_key}': '{original_date}' -> '{formatted_date}'")
                     combined_data[date_key] = formatted_date
                elif not formatted_date:
                     print(f"  -> Warning: Could not re-format date '{date_key}': '{original_date}'")

        if not combined_data.get("pob_state_province") and not combined_data.get("pob_state_province_na"):
            print("  -> Setting pob_state_province_na to True as state/province is missing.")
            combined_data["pob_state_province_na"] = True
        elif combined_data.get("pob_state_province") and combined_data.get("pob_state_province_na"):
             print("  -> Setting pob_state_province_na to False as state/province is present.")
             combined_data["pob_state_province_na"] = False


        # --- End Mapping ---
        print("Final combined data prepared for response.")
        return jsonify(combined_data)

    except Exception as e:
        print(f"Error occurred in /api/getFormData processing: {e}")
        print(traceback.format_exc())
        return jsonify({"message": "An error occurred while processing the form data request."}), 500


# --- CORS Handling ---
# (CORS setup remains the same)
YOUR_EXTENSION_ID = "YOUR_EXTENSION_ID_HERE" # Replace placeholder if needed
if YOUR_EXTENSION_ID == "YOUR_EXTENSION_ID_HERE":
    print("\n*********************************************************************")
    print("WARNING: CHROME_EXTENSION_ID is not set in environment variables.")
    print("CORS will likely fail for the Chrome Extension.")
    print("Set the CHROME_EXTENSION_ID environment variable or update the code.")
    print("*********************************************************************\n")
cors_origin = f"chrome-extension://{YOUR_EXTENSION_ID}"
CORS(app, resources={r"/api/*": {"origins": cors_origin}})
print(f"CORS configured for API routes. Allowing origin: {cors_origin}")
# --- Server Start ---
if __name__ == '__main__':
    print("Starting Flask development server...")

    # Final checks for critical components before starting
    if mongo_client is None:
        print("FATAL: MongoDB connection failed during startup. Server cannot start.")
        exit(1) # Exit if DB connection failed

    # Print warnings for missing optional dependencies
    print("\n--- Optional Dependency Status ---")
    if docx is None: print("WARN: python-docx library not found. .docx file processing will fail.")
    else: print("OK: python-docx library found.")
    if PaddleOCR is None: print("WARN: paddleocr library not found. Image OCR processing will fail.")
    else: print("OK: paddleocr library found.")
    if genai is None: print("WARN: google-generativeai library not found. Gemini analysis will fail.")
    else: print("OK: google-generativeai library found.")
    print("---------------------------------\n")

    # Run the Flask app
    app.run(debug=True, host='0.0.0.0', port=5000)